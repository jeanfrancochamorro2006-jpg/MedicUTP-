# -*- coding: utf-8 -*-
"""
Generador del Informe del Proyecto MEDIC-UTP en formato Word (.docx).

Uso:
    python generar_informe.py

Genera el archivo 'Informe_MedicUTP.docx' en la misma carpeta, listo para
abrir/editar en Microsoft Word. Todo el cuerpo del documento usa
Times New Roman 11 pt y una paleta estrictamente en blanco y negro
(sin colores), según el formato solicitado por el curso.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FUENTE = "Times New Roman"
TAMANIO_BODY = 11
TAMANIO_CODE = 9.5
NEGRO = RGBColor(0x00, 0x00, 0x00)

INTEGRANTES = [
    "Integrante 1: ____________________________________",
    "Integrante 2: ____________________________________",
    "Integrante 3: ____________________________________",
    "Integrante 4: ____________________________________",
    "Integrante 5: ____________________________________",
    "Integrante 6: ____________________________________",
]
DOCENTE = "____________________________________"


# =====================================================================
# UTILIDADES DE FORMATO
# =====================================================================

def fijar_fuente_normal(doc):
    """Configura Times New Roman 11 negro como estilo Normal por defecto."""
    estilo = doc.styles["Normal"]
    estilo.font.name = FUENTE
    estilo.font.size = Pt(TAMANIO_BODY)
    estilo.font.color.rgb = NEGRO
    rpr = estilo.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FUENTE)

    # Los estilos de encabezado de Word traen color de tema (azul) por
    # defecto: se fuerzan a negro para que todo el documento sea B/N.
    for nombre_estilo in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            doc.styles[nombre_estilo].font.color.rgb = NEGRO
        except KeyError:
            pass


def parrafo(doc, texto, negrita=False, cursiva=False, tamanio=TAMANIO_BODY,
            alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY, espacio_despues=6):
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.space_after = Pt(espacio_despues)
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(tamanio)
    run.font.color.rgb = NEGRO
    run.bold = negrita
    run.italic = cursiva
    return p


def titulo_seccion(doc, texto, nivel=1):
    h = doc.add_heading(level=nivel)
    run = h.add_run(texto)
    run.font.name = FUENTE
    run.font.color.rgb = NEGRO
    run.font.size = Pt(16 if nivel == 1 else (13 if nivel == 2 else 12))
    run.bold = True
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(8)
    return h


def vineta(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANIO_BODY)
    run.font.color.rgb = NEGRO
    p.paragraph_format.space_after = Pt(2)
    return p


def bloque_codigo(doc, lineas):
    """Inserta un bloque de texto monoespaciado (caja con borde negro, sin relleno de color)."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    celda = tabla.rows[0].cells[0]
    celda.paragraphs[0].text = ""
    for i, linea_texto in enumerate(lineas):
        p = celda.paragraphs[0] if i == 0 else celda.add_paragraph()
        run = p.add_run(linea_texto if linea_texto else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(TAMANIO_CODE)
        run.font.color.rgb = NEGRO
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tabla_simple(doc, encabezados, filas, anchos_cm=None):
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tabla.rows[0].cells
    for i, texto in enumerate(encabezados):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(texto)
        run.font.name = FUENTE
        run.font.size = Pt(TAMANIO_BODY)
        run.font.color.rgb = NEGRO
        run.bold = True
    for fila in filas:
        celdas = tabla.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run(str(valor))
            run.font.name = FUENTE
            run.font.size = Pt(TAMANIO_BODY)
            run.font.color.rgb = NEGRO
    if anchos_cm:
        for i, ancho in enumerate(anchos_cm):
            for row in tabla.rows:
                row.cells[i].width = Cm(ancho)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tabla


def espacio_para_imagen(doc, etiqueta, alto_cm=4.5, ancho_cm=8):
    """
    Reserva un recuadro con borde negro donde el usuario puede insertar una
    imagen manualmente en Word (Insertar > Imágenes), sin usar color de relleno.
    """
    tabla = doc.add_table(rows=1, cols=1)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    celda = tabla.rows[0].cells[0]
    celda.width = Cm(ancho_cm)
    tr = tabla.rows[0]
    tr.height = Cm(alto_cm)
    celda.vertical_alignment = 1  # centrado vertical
    p = celda.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(etiqueta)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANIO_BODY)
    run.font.color.rgb = NEGRO
    run.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def salto_pagina(doc):
    doc.add_page_break()


# =====================================================================
# PORTADA
# =====================================================================

def construir_portada(doc):
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD TECNOLÓGICA DEL PERÚ")
    run.font.name = FUENTE
    run.font.size = Pt(16)
    run.font.color.rgb = NEGRO
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Facultad de Ingeniería / Escuela de Ingeniería de Sistemas")
    run.font.name = FUENTE
    run.font.size = Pt(12)
    run.font.color.rgb = NEGRO

    doc.add_paragraph()

    # Espacio reservado para el logotipo de MEDIC-UTP
    espacio_para_imagen(doc, "[ Espacio reservado para el logo de MEDIC-UTP ]",
                         alto_cm=4, ancho_cm=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME DEL PROYECTO")
    run.font.name = FUENTE
    run.font.size = Pt(14)
    run.font.color.rgb = NEGRO
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MEDIC-UTP")
    run.font.name = FUENTE
    run.font.size = Pt(22)
    run.font.color.rgb = NEGRO
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Citas Médicas y Facturación en Consola")
    run.font.name = FUENTE
    run.font.size = Pt(13)
    run.font.color.rgb = NEGRO
    run.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrantes:")
    run.font.name = FUENTE
    run.font.size = Pt(TAMANIO_BODY)
    run.font.color.rgb = NEGRO
    run.bold = True

    for texto in INTEGRANTES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.name = FUENTE
        run.font.size = Pt(TAMANIO_BODY)
        run.font.color.rgb = NEGRO

    doc.add_paragraph()

    datos = [
        ("Curso", "Lenguajes de Programación / Programación Orientada a Objetos"),
        ("Docente", DOCENTE),
        ("Sección", "____________________________"),
        ("Fecha", "14 de julio de 2026"),
        ("Ciudad", "Lima, Perú"),
    ]
    for etiqueta, valor in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{etiqueta}: ")
        run.font.name = FUENTE
        run.font.size = Pt(TAMANIO_BODY)
        run.font.color.rgb = NEGRO
        run.bold = True
        run2 = p.add_run(valor)
        run2.font.name = FUENTE
        run2.font.size = Pt(TAMANIO_BODY)
        run2.font.color.rgb = NEGRO

    salto_pagina(doc)


# =====================================================================
# 1. INTRODUCCIÓN
# =====================================================================

def construir_introduccion(doc):
    titulo_seccion(doc, "1. Introducción")
    parrafo(doc,
        "MEDIC-UTP es un sistema de gestión clínica desarrollado íntegramente en "
        "Python y ejecutado por consola. Su propósito es digitalizar el flujo de "
        "atención de una clínica en cuatro etapas secuenciales: registro de "
        "pacientes y médicos, agendamiento de citas médicas, generación de "
        "comprobantes de facturación (boletas y facturas, con descarga en PDF) "
        "y análisis de los datos generados por el propio sistema. El proyecto se "
        "desarrolla como parte del curso de Programación Orientada a Objetos "
        "(POO) de la Universidad Tecnológica del Perú, y busca aplicar de forma "
        "práctica los conceptos de clases, herencia, encapsulamiento, "
        "polimorfismo, programación funcional y análisis de datos con librerías "
        "especializadas."
    )
    parrafo(doc,
        "El sistema no cuenta con interfaz gráfica: toda la interacción con el "
        "usuario se realiza mediante menús numerados en la terminal, validación de "
        "entradas con manejo de excepciones (try/except) y tablas de texto "
        "formateadas para la presentación de resultados."
    )


# =====================================================================
# 2. OBJETIVOS
# =====================================================================

def construir_objetivos(doc):
    titulo_seccion(doc, "2. Objetivos")
    titulo_seccion(doc, "2.1 Objetivo general", nivel=2)
    parrafo(doc,
        "Desarrollar un sistema de consola en Python que permita gestionar el "
        "registro de pacientes y médicos, el agendamiento de citas, la "
        "facturación de atenciones médicas y el análisis estadístico de la "
        "información generada, aplicando los principios de la Programación "
        "Orientada a Objetos, la programación funcional y buenas prácticas de "
        "diseño de software."
    )
    titulo_seccion(doc, "2.2 Objetivos específicos", nivel=2)
    for txt in [
        "Modelar las entidades del dominio (Persona, Paciente, Médico, Cita, "
        "Factura) mediante clases con herencia simple y múltiple.",
        "Aplicar polimorfismo por sobrecarga y por sobreescritura de métodos "
        "en la búsqueda y presentación de información.",
        "Implementar validación robusta de datos de entrada mediante manejo "
        "de excepciones.",
        "Persistir la información del sistema en un archivo JSON para "
        "mantener el estado entre ejecuciones.",
        "Organizar el sistema en módulos independientes por responsabilidad "
        "(registro, agendamiento, facturación, persistencia, análisis).",
        "Analizar los datos registrados con pandas y numpy, y generar "
        "reportes exportables (CSV) y gráficos (matplotlib).",
        "Verificar el correcto funcionamiento de las reglas de negocio "
        "críticas mediante pruebas unitarias automatizadas.",
    ]:
        vineta(doc, txt)


# =====================================================================
# 3. PARADIGMAS Y CONCEPTOS DE PROGRAMACIÓN APLICADOS
# =====================================================================

def construir_paradigmas(doc):
    titulo_seccion(doc, "3. Paradigmas y Conceptos de Programación Aplicados")
    parrafo(doc,
        "Esta sección documenta, en el orden solicitado por el curso, cada uno "
        "de los conceptos y paradigmas de programación presentes en MEDIC-UTP."
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.1 Métodos de tipo procedimental y funciones para entrada de datos", nivel=2)
    parrafo(doc,
        "El módulo modulos/validaciones.py concentra la programación "
        "procedimental del sistema: funciones independientes, sin estado "
        "propio, que reciben un mensaje, validan la entrada del usuario con "
        "try/except y devuelven un valor ya depurado. Estas funciones no "
        "pertenecen a ninguna clase: se invocan directamente desde los menús."
    )
    bloque_codigo(doc, [
        "def leer_entero(mensaje, minimo=None, maximo=None):",
        '    """Lee un entero controlando errores con try/except."""',
        "    while True:",
        "        try:",
        "            valor = int(input(f\"  {mensaje}\"))",
        "            if minimo is not None and valor < minimo:",
        "                mensaje_error(f\"El valor mínimo es {minimo}.\")",
        "                continue",
        "            return valor",
        "        except ValueError:",
        "            mensaje_error(\"Ingrese únicamente números enteros.\")",
    ])
    parrafo(doc,
        "Funciones equivalentes (leer_cadena, leer_dni, leer_fecha, "
        "elegir_de_tupla) validan texto, DNI, fechas y opciones de menú. "
        "Sobre esa base procedimental se construyen la herencia (simple en "
        "Medico, múltiple en Paciente) y el polimorfismo (sobrecarga en "
        "BuscadorClinico y sobreescritura en obtener_resumen()), que se "
        "profundizan en los apartados 3.7 y 3.8."
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.2 Herencia simple y herencia múltiple (resumen)", nivel=2)
    tabla_simple(doc,
        ["Tipo de herencia", "Clases involucradas", "Mecanismo"],
        [
            ("Simple", "Persona → Medico", "super().__init__(dni, nombre)"),
            ("Múltiple", "Persona + InformacionContacto → Paciente",
             "Persona.__init__(...) e InformacionContacto.__init__(...) explícitos"),
        ],
        anchos_cm=[3, 6, 7],
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.3 Polimorfismo (resumen)", nivel=2)
    parrafo(doc,
        "MEDIC-UTP implementa las dos formas clásicas de polimorfismo: por "
        "SOBREESCRITURA, donde Persona, Paciente, Medico y Factura redefinen "
        "obtener_resumen() con su propio formato; y por SOBRECARGA simulada, "
        "donde BuscadorClinico.buscar() cambia de comportamiento según los "
        "parámetros opcionales que recibe (dni, nombre, ambos, o ninguno)."
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.4 Listas, tuplas, colecciones, funciones y módulos", nivel=2)
    tabla_simple(doc,
        ["Estructura", "Uso en el sistema"],
        [
            ("Listas", "pacientes, medicos, citas, facturas — colecciones principales en memoria (modulos/estado.py)."),
            ("Tuplas", "Catálogos inmutables: ESPECIALIDADES_VALIDAS, TIPOS_SEGURO, ESTADOS_CITA, HORAS_ATENCION."),
            ("Colecciones (set)", "dni_pacientes, dni_medicos y slots_ocupados para validar duplicados en O(1)."),
            ("Colecciones (dict)", "indice_especialidad (especialidad → médicos) y acciones (menú → función)."),
            ("Funciones", "Funciones puras reutilizables en modulos/utilidades.py y modulos/ui_consola.py."),
            ("Módulos", "El sistema se organiza en 11 módulos independientes dentro del paquete modulos/."),
        ],
        anchos_cm=[4, 12],
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.5 Programación Orientada a Objetos (POO)", nivel=2)
    parrafo(doc,
        "El dominio del negocio (personas, citas y comprobantes) se modela "
        "íntegramente con clases encapsuladas: todos los atributos son "
        "protegidos (prefijo _) y se accede a ellos únicamente a través de "
        "métodos getter/setter, evitando el acceso directo desde fuera de la "
        "clase. Esto permite validar datos antes de modificarlos (por "
        "ejemplo, InformacionContacto.establecer_telefono() rechaza valores "
        "no numéricos)."
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.6 Clases, objetos y constructores", nivel=2)
    parrafo(doc,
        "Cada entidad del dominio es una clase con su propio constructor "
        "__init__, responsable de inicializar el estado del objeto:"
    )
    bloque_codigo(doc, [
        "class Persona:",
        '    """Clase base (Superclase) que contiene los atributos comunes."""',
        "    def __init__(self, dni, nombre):",
        "        self._dni = dni",
        "        self._nombre = nombre",
        "",
        "    def obtener_dni(self):",
        "        return self._dni",
    ])
    parrafo(doc,
        "A partir de esta clase se crean objetos concretos, por ejemplo "
        "Paciente(\"12345678\", \"Ana Torres\", 30, \"987654321\", \"SIS\"), "
        "cada uno con su propio estado independiente en memoria."
    )

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.7 Herencia simple y múltiple (detalle)", nivel=2)
    parrafo(doc, "Herencia simple — Medico hereda directamente de Persona usando super():")
    bloque_codigo(doc, [
        "class Medico(Persona):",
        '    """Medico demuestra HERENCIA SIMPLE."""',
        "    def __init__(self, dni, nombre, especialidad, consultorio, precio_consulta=0.0):",
        "        super().__init__(dni, nombre)",
        "        self._especialidad = especialidad",
        "        self._consultorio = consultorio",
        "        self._precio_consulta = precio_consulta",
    ])
    parrafo(doc, "Herencia múltiple — Paciente hereda de Persona e InformacionContacto:")
    bloque_codigo(doc, [
        "class Paciente(Persona, InformacionContacto):",
        '    """Paciente demuestra HERENCIA MÚLTIPLE."""',
        "    def __init__(self, dni, nombre, edad, telefono, seguro=\"Particular\"):",
        "        Persona.__init__(self, dni, nombre)",
        "        InformacionContacto.__init__(self, telefono)",
        "        self._edad = edad",
        "        self._seguro = seguro",
    ])

    # -----------------------------------------------------------------
    titulo_seccion(doc, "3.8 Polimorfismo — sobrecarga de métodos", nivel=2)
    parrafo(doc,
        "Python no soporta sobrecarga real de métodos (no admite dos métodos "
        "con el mismo nombre y distinta firma), por lo que se simula con "
        "parámetros opcionales de valor por defecto. BuscadorClinico.buscar() "
        "cambia su comportamiento según qué argumentos recibe, comportándose "
        "como si tuviera cuatro firmas distintas:"
    )
    bloque_codigo(doc, [
        "class BuscadorClinico:",
        "    def buscar(self, coleccion, dni=None, nombre=None):",
        "        if dni is not None and nombre is not None:",
        "            return [e for e in coleccion",
        "                    if e.obtener_dni() == dni",
        "                    and nombre.lower() in e.obtener_nombre().lower()]",
        "        elif dni is not None:",
        "            return [e for e in coleccion if e.obtener_dni() == dni]",
        "        elif nombre is not None:",
        "            return [e for e in coleccion",
        "                    if nombre.lower() in e.obtener_nombre().lower()]",
        "        else:",
        "            return coleccion",
    ])
    parrafo(doc,
        "El polimorfismo por sobreescritura se aplica en paralelo: la función "
        "imprimir_resumen_entidad(entidad) llama a entidad.obtener_resumen() "
        "sin conocer la clase concreta del objeto; el método ejecutado varía "
        "según si la entidad es un Paciente, un Medico o una Factura."
    )


# =====================================================================
# 4. DISEÑO DEL SISTEMA
# =====================================================================

def construir_diseno(doc):
    titulo_seccion(doc, "4. Diseño del Sistema")

    titulo_seccion(doc, "4.1 Arquitectura y estructura de carpetas", nivel=2)
    parrafo(doc,
        "El proyecto separa las entidades del dominio, la lógica de cada "
        "paso del flujo y la interacción con el usuario en paquetes "
        "independientes, organizándose de la siguiente manera:"
    )
    bloque_codigo(doc, [
        "MedicUTP-/",
        "├── main.py                    # Orquesta el menú principal",
        "├── requirements.txt",
        "├── db.json                    # Persistencia local (no versionado en git)",
        "├── clases/",
        "│   ├── persona.py             # Clase base Persona",
        "│   ├── contacto.py            # InformacionContacto (herencia múltiple)",
        "│   ├── paciente.py            # Paciente(Persona, InformacionContacto)",
        "│   ├── medico.py              # Medico(Persona)",
        "│   ├── cita.py                # Cita (composición de objetos)",
        "│   └── factura.py             # Factura (cálculo de montos y descuentos)",
        "├── polimorfismo/",
        "│   ├── sobrecarga.py          # BuscadorClinico (overloading simulado)",
        "│   └── sobreescritura.py      # imprimir_resumen_entidad (overriding)",
        "├── modulos/",
        "│   ├── estado.py              # Colecciones y catálogos compartidos",
        "│   ├── ui_consola.py          # Diseño de pantallas",
        "│   ├── validaciones.py        # Entrada/salida validada",
        "│   ├── utilidades.py          # Búsquedas y RUC (funciones puras)",
        "│   ├── persistencia.py        # Guardado/carga en db.json",
        "│   ├── registro.py",
        "│   ├── agendamiento.py",
        "│   ├── facturacion.py",
        "│   ├── exportar_pdf.py        # Descarga de comprobantes en PDF",
        "│   ├── analisis_datos.py      # pandas / numpy",
        "│   └── visualizacion.py       # Gráficos matplotlib",
        "└── tests/                     # Pruebas unitarias (unittest)",
    ])

    titulo_seccion(doc, "4.2 Modelo de clases", nivel=2)
    tabla_simple(doc,
        ["Clase", "Relación", "Responsabilidad"],
        [
            ("Persona", "Superclase", "DNI y nombre; define obtener_resumen() base."),
            ("InformacionContacto", "Superclase secundaria", "Encapsula el teléfono de contacto."),
            ("Paciente", "Persona + InformacionContacto (herencia múltiple)", "Edad y tipo de seguro."),
            ("Medico", "Persona (herencia simple, super())", "Especialidad, consultorio y precio de consulta."),
            ("Cita", "Composición de Paciente y Medico", "Fecha, hora, motivo y estado de la atención."),
            ("Factura", "Composición de Cita", "Calcula monto, descuento por seguro y total a pagar."),
        ],
        anchos_cm=[3, 5, 8],
    )

    titulo_seccion(doc, "4.3 Persistencia de datos", nivel=2)
    parrafo(doc,
        "Toda la información se guarda en un archivo db.json mediante las "
        "funciones guardar_base_datos() y cargar_base_datos() de "
        "modulos/persistencia.py, utilizando el módulo estándar json. El "
        "archivo se excluye del control de versiones (.gitignore) porque su "
        "contenido cambia en cada ejecución y no forma parte del código "
        "fuente del proyecto."
    )


# =====================================================================
# 5. MÓDULOS FUNCIONALES
# =====================================================================

def construir_modulos(doc):
    titulo_seccion(doc, "5. Descripción de Módulos Funcionales")

    titulo_seccion(doc, "5.1 Módulo de Registro", nivel=2)
    parrafo(doc,
        "Permite registrar y listar pacientes y médicos, validando DNI "
        "(8 dígitos), edad (0-120) y evitando duplicados mediante conjuntos. "
        "Incluye búsqueda por seguro (filter + lambda), búsqueda por "
        "especialidad (índice invertido con diccionario) y estadísticas "
        "(edad promedio, mínima, máxima y distribución por seguro)."
    )

    titulo_seccion(doc, "5.2 Módulo de Agendamiento", nivel=2)
    parrafo(doc,
        "Permite agendar una cita entre un paciente y un médico registrados, "
        "validando fecha (no anterior a hoy), hora (elegida de un catálogo "
        "fijo) y evitando doble reserva del mismo médico en el mismo horario "
        "mediante el conjunto slots_ocupados. También permite listar, "
        "filtrar y cambiar el estado de una cita (Pendiente, Completada, "
        "Cancelada)."
    )

    titulo_seccion(doc, "5.3 Módulo de Facturación", nivel=2)
    parrafo(doc,
        "Genera boletas o facturas únicamente para citas en estado "
        "'Completada'. El descuento se calcula automáticamente según el tipo "
        "de seguro del paciente (ESSALUD 30%, SIS 20%, Particular 0%, Otro "
        "10%). Para facturas, el sistema genera un RUC 10 válido a partir del "
        "DNI del paciente aplicando el algoritmo oficial de dígito "
        "verificador módulo 11 de SUNAT."
    )
    parrafo(doc,
        "El comprobante se construye línea por línea con "
        "construir_lineas_boleta(), que se reutiliza tanto para imprimirlo "
        "en consola como para exportarlo a PDF con "
        "modulos/exportar_pdf.py (reportlab), garantizando que el archivo "
        "descargado sea idéntico al que se muestra en pantalla. El PDF se "
        "guarda automáticamente en reportes/comprobantes/ al generar el "
        "comprobante, y también puede descargarse luego desde el menú "
        "'Descargar Comprobante en PDF' indicando su ID."
    )

    titulo_seccion(doc, "5.4 Módulo de Análisis de Datos (pandas / numpy / matplotlib)", nivel=2)
    parrafo(doc,
        "El módulo modulos/analisis_datos.py transforma las listas de "
        "objetos en memoria en DataFrames de pandas (df_pacientes, "
        "df_medicos, df_citas, df_facturas) para realizar limpieza, "
        "agregación y análisis estadístico:"
    )
    for txt in [
        "Análisis de pacientes: describe() de la edad, media/mediana/"
        "desviación estándar/percentiles con numpy, y distribución por "
        "tipo de seguro con groupby().",
        "Análisis de facturación: ingreso total, ticket promedio e ingresos "
        "por especialidad (groupby + agg), calculados con pandas y "
        "verificados de forma cruzada con functools.reduce() sobre los "
        "objetos Factura (refuerzo del paradigma funcional).",
        "Análisis de agendamiento: citas por estado, por mes y médicos con "
        "más citas asignadas (value_counts()).",
        "Exportación de reportes a CSV (pandas.to_csv) y generación de "
        "gráficos con matplotlib (barras, histograma y gráfico circular) "
        "guardados como PNG en la carpeta reportes/.",
    ]:
        vineta(doc, txt)


# =====================================================================
# 6. PRUEBAS REALIZADAS
# =====================================================================

def construir_pruebas(doc):
    titulo_seccion(doc, "6. Pruebas Realizadas")

    titulo_seccion(doc, "6.1 Pruebas unitarias automatizadas (unittest)", nivel=2)
    parrafo(doc,
        "El paquete tests/ contiene 29 pruebas unitarias ejecutadas con el "
        "módulo estándar unittest, agrupadas en cuatro archivos:"
    )
    tabla_simple(doc,
        ["Archivo", "Qué verifica", "Pruebas"],
        [
            ("test_factura.py", "Cálculo de descuento por tipo de seguro y algoritmo del RUC 10 (dígito verificador módulo 11).", "13"),
            ("test_herencia_poo.py", "Herencia simple, herencia múltiple, encapsulamiento y sobreescritura de obtener_resumen().", "9"),
            ("test_polimorfismo_sobrecarga.py", "Las cuatro firmas simuladas de BuscadorClinico.buscar() y la función polimórfica imprimir_resumen_entidad().", "6"),
            ("test_analisis_datos.py", "Construcción de DataFrames y coincidencia entre pandas.sum() y functools.reduce().", "3"),
        ],
        anchos_cm=[5, 9, 2],
    )
    parrafo(doc,
        "Comando de ejecución: python -m unittest discover -s tests -v — "
        "resultado obtenido: 29 pruebas ejecutadas, 29 exitosas (OK)."
    )

    titulo_seccion(doc, "6.2 Pruebas manuales de extremo a extremo", nivel=2)
    parrafo(doc,
        "Adicionalmente se recorrió manualmente el flujo completo REGISTRO → "
        "AGENDAMIENTO → FACTURACIÓN → ANÁLISIS DE DATOS, verificando tanto "
        "los casos válidos como los casos de error controlado."
    )
    tabla_simple(doc,
        ["N°", "Caso de prueba", "Resultado esperado", "Resultado obtenido"],
        [
            ("1", "Registrar paciente con DNI de 7 dígitos", "Mensaje de error y nueva solicitud de DNI", "Correcto"),
            ("2", "Registrar paciente con DNI ya existente", "Rechaza el registro duplicado", "Correcto"),
            ("3", "Agendar cita con médico/fecha/hora ya ocupados", "Rechaza el agendamiento (slot ocupado)", "Correcto"),
            ("4", "Agendar cita con fecha anterior a hoy", "Rechaza la fecha e informa el error", "Correcto"),
            ("5", "Generar factura de una cita no 'Completada'", "Bloquea la emisión del comprobante", "Correcto"),
            ("6", "Generar factura duplicada para la misma cita", "Bloquea e informa el comprobante previo", "Correcto"),
            ("7", "Descargar un comprobante ya emitido en PDF", "Genera el PDF en reportes/comprobantes/", "Correcto"),
            ("8", "Analizar pacientes/facturación/citas con pandas", "Muestra estadísticas y agregaciones correctas", "Correcto"),
            ("9", "Generar gráficos con matplotlib", "Guarda 4 imágenes PNG en reportes/", "Correcto"),
            ("10", "Reiniciar el programa tras registrar datos", "Los datos persisten al recargar db.json", "Correcto"),
        ],
        anchos_cm=[1, 6, 6, 3],
    )


# =====================================================================
# 7. RESULTADOS Y VISUALIZACIÓN
# =====================================================================

def construir_resultados(doc):
    titulo_seccion(doc, "7. Resultados y Visualización")
    parrafo(doc,
        "A continuación se muestra un extracto real de la ejecución del "
        "sistema por consola, correspondiente a la pantalla de bienvenida y "
        "al resumen general que se presenta antes del menú principal:"
    )
    bloque_codigo(doc, [
        "============================================================",
        "                       MEDIC-UTP  v1.0",
        "============================================================",
        "           Sistema de Citas Medicas y Facturacion",
        "============================================================",
        "",
        "============================================================",
        "                    RESUMEN DEL SISTEMA",
        "============================================================",
        "  Fecha y hora             : 14/07/2026  10:40",
        "  Pacientes registrados    : 2",
        "  Medicos registrados      : 3",
        "  Especialidades activas   : 3",
        "  Citas agendadas          : 2",
        "  Citas pendientes         : 0",
        "  Comprobantes emitidos    : 2",
    ])

    titulo_seccion(doc, "7.1 Gráficos generados (matplotlib)", nivel=2)
    parrafo(doc,
        "Al ejecutar la opción 'Generar Gráficos' del módulo de Análisis de "
        "Datos, el sistema guarda automáticamente las siguientes imágenes en "
        "la carpeta reportes/: pacientes_por_seguro.png (barras), "
        "edad_pacientes.png (histograma), ingresos_por_especialidad.png "
        "(barras) y citas_por_estado.png (circular). Se reserva el siguiente "
        "espacio para insertar una captura de uno de estos gráficos:"
    )
    espacio_para_imagen(doc, "[ Espacio para insertar un gráfico de reportes/ (ej. ingresos_por_especialidad.png) ]",
                         alto_cm=7, ancho_cm=12)

    parrafo(doc,
        "El sistema fue ejecutado de extremo a extremo (registro, "
        "agendamiento, facturación con descarga en PDF y análisis de datos) "
        "sin errores de ejecución, y con la totalidad de las pruebas "
        "unitarias en estado exitoso."
    )


# =====================================================================
# 8. CÓDIGO COMENTADO
# =====================================================================

def construir_codigo_comentado(doc):
    titulo_seccion(doc, "8. Fragmento de Código Comentado")
    parrafo(doc,
        "Se incluye como ejemplo el método que calcula el total de una "
        "factura, ilustrando el uso de composición de objetos, condicionales "
        "y documentación en línea (docstring):"
    )
    bloque_codigo(doc, [
        "def calcular_total(self):",
        '    """Calcula el monto_consulta, descuento y total de la factura."""',
        "    medico = self._cita.obtener_medico()",
        "    self._monto_consulta = float(medico.obtener_precio_consulta())",
        "",
        "    paciente = self._cita.obtener_paciente()",
        "    seguro = paciente.obtener_seguro()",
        "",
        "    # Descuentos: ESSALUD 30% | SIS 20% | Particular 0% | Otro 10%",
        "    if seguro == \"ESSALUD\":",
        "        porcentaje = 0.30",
        "    elif seguro == \"SIS\":",
        "        porcentaje = 0.20",
        "    elif seguro == \"Particular\":",
        "        porcentaje = 0.0",
        "    else:",
        "        porcentaje = 0.10",
        "",
        "    self._descuento = self._monto_consulta * porcentaje",
        "    self._total = self._monto_consulta - self._descuento",
    ])
    parrafo(doc,
        "Y el cálculo del ingreso total en el módulo de análisis, resuelto "
        "de dos formas equivalentes para reforzar el paradigma funcional "
        "(pandas y functools.reduce):"
    )
    bloque_codigo(doc, [
        "total_pandas = df[\"total\"].sum()",
        "total_reduce = reduce(lambda acc, f: acc + f.obtener_total(),",
        "                      estado.facturas, 0.0)",
    ])


# =====================================================================
# 9. CONCLUSIONES
# =====================================================================

def construir_conclusiones(doc):
    titulo_seccion(doc, "9. Conclusiones")
    for txt in [
        "El sistema MEDIC-UTP cumple con el flujo funcional completo "
        "planteado (Registro, Agendamiento, Facturación y Análisis de "
        "Datos) operando de forma estable por consola.",
        "El uso de herencia simple y múltiple, encapsulamiento y "
        "polimorfismo permitió modelar el dominio clínico de forma clara y "
        "extensible.",
        "La combinación de listas, tuplas, conjuntos y diccionarios "
        "optimizó las búsquedas y evitó datos duplicados o inconsistentes.",
        "La incorporación de pandas y numpy permitió transformar las "
        "colecciones en memoria en información de valor (ingresos por "
        "especialidad, distribución de edades, citas por mes), reforzando "
        "el enfoque multiparadigma del proyecto.",
        "La modularización del sistema en el paquete modulos/ mejoró la "
        "legibilidad y mantenibilidad del código frente a la versión inicial "
        "concentrada en un único archivo.",
        "Las 29 pruebas unitarias automatizadas respaldan la corrección de "
        "las reglas de negocio críticas (cálculo de descuentos, generación "
        "del RUC, herencia y polimorfismo).",
    ]:
        vineta(doc, txt)


# =====================================================================
# 10. FUENTES
# =====================================================================

def construir_fuentes(doc):
    titulo_seccion(doc, "10. Fuentes")
    for txt in [
        "Python Software Foundation. (2025). Python 3 Documentation. "
        "https://docs.python.org/3/",
        "The pandas development team. pandas documentation. "
        "https://pandas.pydata.org/docs/",
        "NumPy developers. NumPy documentation. https://numpy.org/doc/",
        "Matplotlib developers. Matplotlib documentation. https://matplotlib.org/",
        "SUNAT. Algoritmo de cálculo del dígito verificador del RUC "
        "(módulo 11). https://www.sunat.gob.pe/",
        "Material de clase del curso de Programación Orientada a Objetos, "
        "Universidad Tecnológica del Perú (2026).",
    ]:
        vineta(doc, txt)


# =====================================================================
# MAIN
# =====================================================================

def main():
    doc = Document()
    fijar_fuente_normal(doc)

    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    construir_portada(doc)
    construir_introduccion(doc)
    construir_objetivos(doc)
    construir_paradigmas(doc)
    construir_diseno(doc)
    construir_modulos(doc)
    construir_pruebas(doc)
    construir_resultados(doc)
    construir_codigo_comentado(doc)
    construir_conclusiones(doc)
    construir_fuentes(doc)

    salida = "Informe_MedicUTP.docx"
    doc.save(salida)
    print(f"Informe generado correctamente: {salida}")


if __name__ == "__main__":
    main()
